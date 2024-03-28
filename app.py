import os
from flask import Flask, request, render_template, send_file, send_from_directory, jsonify
from werkzeug.utils import secure_filename
from docx import Document
from docx.oxml.shared import qn
from docxtpl import DocxTemplate
from io import BytesIO

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'docx'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/create_docx')
def create_docx():
    doc = Document()

    doc.add_paragraph('Hello, World!')
    doc.add_paragraph('Bim Bim Bim Bam Bam Bam')

    byte_stream = BytesIO()
    doc.save(byte_stream)

    byte_stream.seek(0)

    return byte_stream
    # return send_file(byte_stream, as_attachment=True, download_name='document.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/hello', methods=['GET'])
def hello():
    return "Hello World"

@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        # Проверяем, что файл был загружен
        if 'file' not in request.files:
            return 'No file part'

        file = request.files['file']

        # Проверяем, что файл имеет допустимое расширение
        if file.filename == '':
            return 'No selected file'

        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)

            # Проверяем и создаем папку uploads, если она не существует
            if not os.path.exists(app.config['UPLOAD_FOLDER']):
                os.makedirs(app.config['UPLOAD_FOLDER'])

            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)

            # Открываем документ Word
            doc = Document(file_path)

            # Получаем список закладок
            bookmarks = doc.part.element.findall(".//" + qn("w:bookmarkStart"))

            # Выводим имена закладок
            bookmarks_list = [bookmark.get(qn("w:name")) for bookmark in bookmarks]

            context = {}

            for i, bookmark in enumerate(bookmarks_list):
                context[f"item{i}"] = bookmark

            # Создаем заполненный шаблон
            filled_template_path = os.path.join(app.config['UPLOAD_FOLDER'], 'filled_template.docx')
            book = DocxTemplate(file_path)
            book.render(context)
            book.save(filled_template_path)

            # Возвращаем заполненный шаблон пользователю для скачивания
            return send_file(filled_template_path, as_attachment=True)

    return render_template('upload.html')


if __name__ == '__main__':
    app.run(debug=True)
